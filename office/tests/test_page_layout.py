from __future__ import annotations

from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION

from inkspan_office import (
    OfficeDocumentError,
    load_schema,
    render_office_document,
    write_office_document,
)
from inkspan_office.page_layout import apply_docx_page_layout

_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_DOC_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)
_PNG_DATA_URI = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9WlXkWQAAAAASUVORK5CYII="
)


def _layout_payload() -> dict[str, object]:
    """Return one valid bounded landscape A4 request used by fidelity tests."""

    return {
        "format": "docx",
        "page_layout": {
            "paper_size": "a4",
            "orientation": "landscape",
            "margins_mm": {"top": 10, "right": 20, "bottom": 30, "left": 40},
        },
        "blocks": [{"type": "paragraph", "text": "Layout contract"}],
    }


def _portrait_letter_layout() -> dict[str, object]:
    """Return one valid Letter portrait layout used to cover alternate dimensions."""

    return {
        "paper_size": "letter",
        "orientation": "portrait",
        "margins_mm": {"top": 0, "right": 5, "bottom": 10, "left": 15},
    }


def test_page_layout_is_part_of_the_machine_readable_docx_contract() -> None:
    schema = load_schema()
    docx_schema = schema["oneOf"][0]

    assert docx_schema["properties"]["page_layout"] == {
        "$ref": "#/$defs/pageLayout"
    }
    layout = schema["$defs"]["pageLayout"]
    assert layout["required"] == ["paper_size", "orientation", "margins_mm"]
    assert layout["additionalProperties"] is False
    assert layout["properties"]["paper_size"]["enum"] == ["a4", "letter"]
    assert layout["properties"]["orientation"]["enum"] == [
        "portrait",
        "landscape",
    ]
    margins = schema["$defs"]["pageMargins"]
    assert margins["required"] == ["top", "right", "bottom", "left"]
    assert margins["additionalProperties"] is False
    for name in margins["required"]:
        assert margins["properties"][name] == {
            "type": "integer",
            "minimum": 0,
            "maximum": 100,
        }


def test_render_docx_applies_bounded_page_layout() -> None:
    rendered = render_office_document(_layout_payload())

    document = Document(BytesIO(rendered.data))
    assert len(document.sections) == 1
    section = document.sections[0]
    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert section.page_width is not None
    assert section.page_height is not None
    assert section.top_margin is not None
    assert section.right_margin is not None
    assert section.bottom_margin is not None
    assert section.left_margin is not None
    assert section.page_width.mm == pytest.approx(297, abs=0.02)
    assert section.page_height.mm == pytest.approx(210, abs=0.02)
    assert section.top_margin.mm == pytest.approx(10, abs=0.02)
    assert section.right_margin.mm == pytest.approx(20, abs=0.02)
    assert section.bottom_margin.mm == pytest.approx(30, abs=0.02)
    assert section.left_margin.mm == pytest.approx(40, abs=0.02)


def test_render_docx_applies_letter_portrait_page_layout() -> None:
    rendered = render_office_document(
        {
            "format": "docx",
            "page_layout": _portrait_letter_layout(),
            "blocks": [],
        }
    )

    section = Document(BytesIO(rendered.data)).sections[0]
    assert section.orientation == WD_ORIENT.PORTRAIT
    assert section.page_width is not None
    assert section.page_height is not None
    assert section.top_margin is not None
    assert section.right_margin is not None
    assert section.bottom_margin is not None
    assert section.left_margin is not None
    assert section.page_width.mm == pytest.approx(215.9, abs=0.02)
    assert section.page_height.mm == pytest.approx(279.4, abs=0.02)
    assert section.top_margin.mm == pytest.approx(0, abs=0.02)
    assert section.right_margin.mm == pytest.approx(5, abs=0.02)
    assert section.bottom_margin.mm == pytest.approx(10, abs=0.02)
    assert section.left_margin.mm == pytest.approx(15, abs=0.02)


def test_render_docx_page_layout_preserves_external_hyperlink_relationship() -> None:
    payload = _layout_payload()
    payload["blocks"] = [
        {
            "type": "rich_paragraph",
            "runs": [
                {
                    "text": "Inkspan report",
                    "bold": True,
                    "href": "https://example.com/report",
                }
            ],
        }
    ]

    rendered = render_office_document(payload)
    with ZipFile(BytesIO(rendered.data)) as package:
        document_root = ElementTree.fromstring(package.read("word/document.xml"))
        relationships_root = ElementTree.fromstring(
            package.read("word/_rels/document.xml.rels")
        )

    hyperlink = document_root.find(f".//{{{_WORD_NS}}}hyperlink")
    assert hyperlink is not None
    relationship_id = hyperlink.attrib[f"{{{_DOC_REL_NS}}}id"]
    assert "".join(hyperlink.itertext()) == "Inkspan report"
    assert hyperlink.find(f".//{{{_WORD_NS}}}b") is not None

    relationship = next(
        relation
        for relation in relationships_root.findall(f"{{{_REL_NS}}}Relationship")
        if relation.attrib.get("Id") == relationship_id
    )
    assert relationship.attrib["Type"] == _HYPERLINK_REL
    assert relationship.attrib["Target"] == "https://example.com/report"
    assert relationship.attrib["TargetMode"] == "External"


def test_render_docx_page_layout_preserves_inline_image_and_alt_text() -> None:
    payload = _layout_payload()
    payload["blocks"] = [
        {
            "type": "image",
            "source": _PNG_DATA_URI,
            "alt_text": "One-pixel layout fidelity fixture",
            "width_px": 96,
        }
    ]

    rendered = render_office_document(payload)
    document = Document(BytesIO(rendered.data))
    assert len(document.inline_shapes) == 1
    assert document.inline_shapes[0].width == 96 * 9_525
    assert document.inline_shapes[0].height == 96 * 9_525

    with ZipFile(BytesIO(rendered.data)) as package:
        image_parts = [name for name in package.namelist() if name.startswith("word/media/")]
        assert len(image_parts) == 1
        document_xml = package.read("word/document.xml").decode("utf-8")
        assert 'descr="One-pixel layout fidelity fixture"' in document_xml


def test_render_docx_page_layout_is_deterministic() -> None:
    first = render_office_document(_layout_payload())
    second = render_office_document(_layout_payload())

    assert first.data == second.data


def test_page_layout_rejects_multiple_docx_sections() -> None:
    document = Document()
    document.add_section(WD_SECTION.NEW_PAGE)
    source = BytesIO()
    document.save(source)

    with pytest.raises(OfficeDocumentError, match="exactly one DOCX section"):
        apply_docx_page_layout(source.getvalue(), _portrait_letter_layout())


def test_invalid_page_layout_never_publishes_partial_output(tmp_path: Path) -> None:
    output = tmp_path / "invalid-layout.docx"
    payload = {
        "format": "docx",
        "page_layout": {
            "paper_size": "a4",
            "orientation": "portrait",
            "margins_mm": {"top": -1, "right": 10, "bottom": 10, "left": 10},
        },
        "blocks": [{"type": "paragraph", "text": "must not publish"}],
    }

    with pytest.raises(OfficeDocumentError):
        write_office_document(payload, output)

    assert not output.exists()


@pytest.mark.parametrize(
    "page_layout",
    [
        None,
        {1: "not-a-string-key"},
        {"paper_size": "a4", "orientation": "portrait"},
        {
            "paper_size": "a4",
            "orientation": "portrait",
            "margins_mm": [],
        },
        {
            "paper_size": 1,
            "orientation": "portrait",
            "margins_mm": {"top": 10, "right": 10, "bottom": 10, "left": 10},
        },
        {
            "paper_size": "A4",
            "orientation": "portrait",
            "margins_mm": {"top": 10, "right": 10, "bottom": 10, "left": 10},
        },
        {
            "paper_size": "letter",
            "orientation": "PORTRAIT",
            "margins_mm": {"top": 10, "right": 10, "bottom": 10, "left": 10},
        },
        {
            "paper_size": "letter",
            "orientation": "portrait",
            "margins_mm": {"top": True, "right": 10, "bottom": 10, "left": 10},
        },
        {
            "paper_size": "letter",
            "orientation": "portrait",
            "margins_mm": {"top": 101, "right": 10, "bottom": 10, "left": 10},
        },
        {
            "paper_size": "letter",
            "orientation": "portrait",
            "margins_mm": {
                "top": 10,
                "right": 10,
                "bottom": 10,
                "left": 10,
                "gutter": 1,
            },
        },
        {
            "paper_size": "letter",
            "orientation": "portrait",
            "margins_mm": {"top": 10, "right": 10, "bottom": 10, "left": 10},
            "columns": 2,
        },
    ],
)
def test_render_docx_rejects_invalid_page_layout(page_layout: object) -> None:
    with pytest.raises(OfficeDocumentError):
        render_office_document(
            {
                "format": "docx",
                "page_layout": page_layout,
                "blocks": [],
            }
        )
