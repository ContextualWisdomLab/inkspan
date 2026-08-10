"""Contract tests for bounded heading alignment in deterministic DOCX output."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from inkspan_office import (
    OfficeDocumentError,
    load_schema,
    render_office_document,
    write_office_document,
)

_ALIGNMENT_VALUES = ("left", "center", "right", "justify")
_OMITTED = object()


def _heading_schema_branch() -> dict[str, object]:
    """Return the public DOCX heading schema branch."""

    schema = load_schema()
    return next(
        branch
        for branch in schema["$defs"]["docxBlock"]["oneOf"]
        if branch.get("properties", {}).get("type", {}).get("const") == "heading"
    )


def _heading_payload(alignment: object = _OMITTED) -> dict[str, object]:
    """Build one DOCX heading request with optional alignment."""

    block: dict[str, object] = {"type": "heading", "level": 2, "text": "Evidence"}
    if alignment is not _OMITTED:
        block["alignment"] = alignment
    return {"format": "docx", "blocks": [block]}


def test_docx_heading_alignment_is_a_bounded_public_contract() -> None:
    """Heading schema must reuse the existing four-value alignment enum."""

    branch = _heading_schema_branch()
    assert branch["properties"]["alignment"] == {
        "type": "string",
        "enum": list(_ALIGNMENT_VALUES),
    }


@pytest.mark.parametrize(
    ("alignment", "expected"),
    [
        ("left", WD_ALIGN_PARAGRAPH.LEFT),
        ("center", WD_ALIGN_PARAGRAPH.CENTER),
        ("right", WD_ALIGN_PARAGRAPH.RIGHT),
        ("justify", WD_ALIGN_PARAGRAPH.JUSTIFY),
    ],
)
def test_docx_heading_alignment_round_trips(
    alignment: str, expected: WD_ALIGN_PARAGRAPH
) -> None:
    """Every accepted heading alignment must round-trip through DOCX."""

    rendered = render_office_document(_heading_payload(alignment))
    paragraph = Document(BytesIO(rendered.data)).paragraphs[0]
    assert paragraph.style.name == "Heading 2"
    assert paragraph.alignment == expected

    with ZipFile(BytesIO(rendered.data)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:jc" in document_xml


def test_docx_heading_omitted_alignment_preserves_inherited_word_state() -> None:
    """Omitted heading alignment must not materialize justification."""

    rendered = render_office_document(_heading_payload())
    paragraph = Document(BytesIO(rendered.data)).paragraphs[0]
    assert paragraph.alignment is None
    with ZipFile(BytesIO(rendered.data)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:jc" not in document_xml


@pytest.mark.parametrize(
    "invalid_alignment",
    ["CENTER", " center", "center ", "distributed", "", 1, True, None],
)
def test_docx_heading_alignment_rejects_invalid_values(
    invalid_alignment: object,
) -> None:
    """Unsupported heading alignment values must fail closed."""

    with pytest.raises(OfficeDocumentError, match=r"blocks\[0\]\.alignment"):
        render_office_document(_heading_payload(invalid_alignment))


def test_docx_heading_alignment_output_is_deterministic() -> None:
    """Explicit heading alignment must preserve deterministic bytes."""

    payload = _heading_payload("center")
    assert render_office_document(payload).data == render_office_document(payload).data


def test_invalid_docx_heading_alignment_never_partially_publishes(tmp_path: Path) -> None:
    """Invalid heading alignment must fail before output publication."""

    destination = tmp_path / "invalid-heading-alignment.docx"
    with pytest.raises(OfficeDocumentError):
        write_office_document(_heading_payload("distributed"), destination)
    assert not destination.exists()
