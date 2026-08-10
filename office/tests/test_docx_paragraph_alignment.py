"""Contract tests for bounded paragraph alignment in deterministic DOCX output."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from inkspan_office import (
    OfficeDocumentError,
    load_schema,
    render_office_document,
    write_office_document,
)

_ALIGNMENT_VALUES = ("left", "center", "right", "justify")
_OMITTED = object()


def _docx_branch(block_type: str) -> dict[str, object]:
    """Return one public DOCX schema branch by its literal block type."""

    schema = load_schema()
    return next(
        branch
        for branch in schema["$defs"]["docxBlock"]["oneOf"]
        if branch.get("properties", {}).get("type", {}).get("const") == block_type
    )


def _plain_payload(alignment: object = _OMITTED) -> dict[str, object]:
    """Build one plain DOCX paragraph request with optional alignment."""

    block: dict[str, object] = {"type": "paragraph", "text": "Evidence"}
    if alignment is not _OMITTED:
        block["alignment"] = alignment
    return {"format": "docx", "blocks": [block]}


def _rich_payload(alignment: object = _OMITTED) -> dict[str, object]:
    """Build one rich DOCX paragraph request with optional alignment."""

    block: dict[str, object] = {
        "type": "rich_paragraph",
        "runs": [{"text": "Rich ", "bold": True}, {"text": "evidence"}],
    }
    if alignment is not _OMITTED:
        block["alignment"] = alignment
    return {"format": "docx", "blocks": [block]}


def test_docx_paragraph_alignment_is_a_bounded_public_contract() -> None:
    """Plain and rich paragraph schemas must expose the same four-value enum."""

    for block_type in ("paragraph", "rich_paragraph"):
        branch = _docx_branch(block_type)
        assert branch["properties"]["alignment"] == {
            "type": "string",
            "enum": list(_ALIGNMENT_VALUES),
        }


@pytest.mark.parametrize(
    ("alignment", "expected"),
    [
        ("left", WD_ALIGN_PARAGRAPH.LEFT),
        ("center", WD_ALIGN_PARAGRAPH.CENTER),
        ("right", WD_ALIGN_PARAGRAPH.RIGHT),
        ("justify", WD_ALIGN_PARAGRAPH.JUSTIFY),
    ],
)
@pytest.mark.parametrize("payload_builder", [_plain_payload, _rich_payload])
def test_docx_paragraph_alignment_round_trips(
    alignment: str,
    expected: WD_ALIGN_PARAGRAPH,
    payload_builder: object,
) -> None:
    """Every accepted alignment must round-trip through the public renderer."""

    payload = payload_builder(alignment)  # type: ignore[operator]
    rendered = render_office_document(payload)
    paragraph = Document(BytesIO(rendered.data)).paragraphs[0]
    assert paragraph.alignment == expected

    with ZipFile(BytesIO(rendered.data)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:jc" in document_xml


def test_docx_omitted_alignment_preserves_inherited_word_state() -> None:
    """Omitting alignment must not materialize a justification property."""

    for payload in (_plain_payload(), _rich_payload()):
        rendered = render_office_document(payload)
        paragraph = Document(BytesIO(rendered.data)).paragraphs[0]
        assert paragraph.alignment is None
        with ZipFile(BytesIO(rendered.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        assert "<w:jc" not in document_xml


@pytest.mark.parametrize(
    "invalid_alignment",
    ["CENTER", " center", "center ", "distributed", "", 1, True, None],
)
@pytest.mark.parametrize("payload_builder", [_plain_payload, _rich_payload])
def test_docx_paragraph_alignment_rejects_invalid_values(
    invalid_alignment: object,
    payload_builder: object,
) -> None:
    """Unsupported or repaired alignment values must fail closed."""

    payload = payload_builder(invalid_alignment)  # type: ignore[operator]
    with pytest.raises(OfficeDocumentError, match=r"blocks\[0\]\.alignment"):
        render_office_document(payload)


def test_docx_paragraph_alignment_output_is_deterministic() -> None:
    """Explicit alignment must preserve byte-identical deterministic output."""

    payload = _rich_payload("justify")
    assert render_office_document(payload).data == render_office_document(payload).data


def test_invalid_docx_alignment_never_partially_publishes(tmp_path: Path) -> None:
    """Alignment validation must fail before any output artifact is published."""

    destination = tmp_path / "invalid-alignment.docx"
    with pytest.raises(OfficeDocumentError):
        write_office_document(_plain_payload("distributed"), destination)
    assert not destination.exists()
