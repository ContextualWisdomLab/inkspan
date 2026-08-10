"""Contract tests for deterministic rich-text runs in DOCX paragraphs."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from inkspan_office import (
    OfficeDocumentError,
    load_schema,
    render_office_document,
    write_office_document,
)


def _rich_payload(runs: list[object]) -> dict[str, object]:
    """Build one minimal DOCX request containing a rich paragraph."""

    return {
        "format": "docx",
        "blocks": [{"type": "rich_paragraph", "runs": runs}],
    }


def test_docx_contract_preserves_explicit_rich_text_runs() -> None:
    """The public schema and renderer must preserve run text and emphasis in order."""

    schema = load_schema()
    rich_branches = [
        branch
        for branch in schema["$defs"]["docxBlock"]["oneOf"]
        if branch.get("properties", {}).get("type", {}).get("const")
        == "rich_paragraph"
    ]
    assert len(rich_branches) == 1

    rendered = render_office_document(
        _rich_payload(
            [
                {"text": "Retention ", "bold": True},
                {"text": "improved", "italic": True},
                {"text": " year over year.", "underline": True},
                {
                    "text": " 검증",
                    "bold": True,
                    "italic": True,
                    "underline": True,
                },
            ]
        )
    )

    document = Document(BytesIO(rendered.data))
    paragraph = document.paragraphs[0]
    assert [run.text for run in paragraph.runs] == [
        "Retention ",
        "improved",
        " year over year.",
        " 검증",
    ]
    assert [(run.bold, run.italic, run.underline) for run in paragraph.runs] == [
        (True, None, None),
        (None, True, None),
        (None, None, True),
        (True, True, True),
    ]


def test_docx_rich_paragraph_preserves_unicode_order_and_false_flags() -> None:
    """Combining, CJK, and bidi text must retain logical order and explicit false flags."""

    rendered = render_office_document(
        _rich_payload(
            [
                {"text": "e\u0301", "bold": False},
                {"text": "漢字", "italic": False},
                {"text": "مرحبا", "underline": False},
            ]
        )
    )

    paragraph = Document(BytesIO(rendered.data)).paragraphs[0]
    assert [run.text for run in paragraph.runs] == ["e\u0301", "漢字", "مرحبا"]
    assert [(run.bold, run.italic, run.underline) for run in paragraph.runs] == [
        (False, None, None),
        (None, False, None),
        (None, None, False),
    ]


def test_docx_rich_paragraph_output_is_deterministic() -> None:
    """The same rich-run request must produce byte-identical canonical OOXML."""

    payload = _rich_payload(
        [
            {"text": "Stable ", "bold": True},
            {"text": "evidence", "italic": True},
        ]
    )
    assert render_office_document(payload).data == render_office_document(payload).data


def test_docx_rich_paragraph_rejects_empty_run_collection() -> None:
    """Runtime validation must reject a rich paragraph that contains no runs."""

    with pytest.raises(
        OfficeDocumentError,
        match=r"blocks\[0\]\.runs must contain at least one run",
    ):
        render_office_document(_rich_payload([]))


def test_docx_rich_paragraph_rejects_runtime_run_overflow() -> None:
    """Runtime validation must retain a finite defense-in-depth run ceiling."""

    with pytest.raises(
        OfficeDocumentError,
        match=r"blocks\[0\]\.runs must contain at most 4096 runs",
    ):
        render_office_document(_rich_payload([{"text": "x"}] * 4097))


@pytest.mark.parametrize(
    ("runs", "message"),
    [
        (["not-an-object"], r"blocks\[0\]\.runs\[0\] must be an object"),
        (
            [{"text": "x", "color": "red"}],
            r"blocks\[0\]\.runs\[0\] has unexpected field: color",
        ),
        (
            [{"text": "x", "bold": 1}],
            r"blocks\[0\]\.runs\[0\]\.bold must be a boolean",
        ),
    ],
)
def test_docx_rich_paragraph_rejects_invalid_run_shapes(
    runs: list[object], message: str
) -> None:
    """Rich runs must remain bounded to object, field, and strict-boolean contracts."""

    with pytest.raises(OfficeDocumentError, match=message):
        render_office_document(_rich_payload(runs))


def test_invalid_rich_paragraph_never_partially_publishes(tmp_path: Path) -> None:
    """Validation must fail before an invalid rich paragraph creates an output file."""

    destination = tmp_path / "invalid-rich.docx"
    with pytest.raises(OfficeDocumentError):
        write_office_document(
            _rich_payload([{"text": "x", "underline": "yes"}]),
            destination,
        )
    assert not destination.exists()
