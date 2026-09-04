from __future__ import annotations

from io import BytesIO

from docx import Document

from inkspan_office import render_office_document


def test_page_layout_preserves_metadata_lists_and_tables() -> None:
    """Prove the layout round trip keeps representative pre-existing DOCX structure."""

    rendered = render_office_document(
        {
            "format": "docx",
            "title": "Layout fidelity report",
            "author": "Inkspan",
            "subject": "Structural round-trip evidence",
            "page_layout": {
                "paper_size": "a4",
                "orientation": "portrait",
                "margins_mm": {"top": 15, "right": 15, "bottom": 15, "left": 15},
            },
            "blocks": [
                {
                    "type": "bullet_list",
                    "ordered": False,
                    "items": ["First finding", "Second finding"],
                },
                {
                    "type": "table",
                    "headers": ["Metric", "Value"],
                    "rows": [["Coverage", "100%"], ["Mode", "deterministic"]],
                },
            ],
        }
    )

    document = Document(BytesIO(rendered.data))
    assert document.core_properties.title == "Layout fidelity report"
    assert document.core_properties.author == "Inkspan"
    assert document.core_properties.subject == "Structural round-trip evidence"

    paragraphs = {paragraph.text: paragraph.style.name for paragraph in document.paragraphs}
    assert paragraphs["First finding"] == "List Bullet"
    assert paragraphs["Second finding"] == "List Bullet"

    assert len(document.tables) == 1
    assert [[cell.text for cell in row.cells] for row in document.tables[0].rows] == [
        ["Metric", "Value"],
        ["Coverage", "100%"],
        ["Mode", "deterministic"],
    ]
